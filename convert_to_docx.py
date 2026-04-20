"""Convert pitch markdown files to Word documents."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_paragraph(doc, text, bold=False, italic=False, align=None, rtl=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    if rtl:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

def md_to_docx(md_path, docx_path, is_rtl=False):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    i = 0
    in_table = False
    in_code = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n').rstrip('\r')
        
        # Skip HTML divs and empty lines in divs
        if line.strip().startswith('<div') or line.strip().startswith('</div>'):
            i += 1
            continue
            
        # Code blocks
        if line.strip().startswith('```'):
            if in_code:
                in_code = False
                i += 1
                continue
            else:
                in_code = True
                i += 1
                continue
        
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # Table rows
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # skip separator rows
            if cells and all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # check if next line is not a table
            if i + 1 >= len(lines) or not lines[i+1].strip().startswith('|'):
                # flush table
                if table_rows:
                    cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=cols)
                    table.style = 'Light Grid Accent 1'
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            table.rows[ri].cells[ci].text = cell_text
                    table_rows = []
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith('### '):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[4:].strip())
            add_heading(doc, text, level=3)
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        elif line.startswith('> '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.style = doc.styles['Quote'] if 'Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(12)
        elif line.startswith('- '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        elif line.strip().startswith(('1.', '2.', '3.', '4.')):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        elif line.strip():
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            add_paragraph(doc, text, rtl=is_rtl)
        
        i += 1

    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

if __name__ == '__main__':
    md_to_docx('pitch.md', 'pitch_ar.docx', is_rtl=True)
    md_to_docx('pitch_en.md', 'pitch_en.docx', is_rtl=False)
    print("Done!")
